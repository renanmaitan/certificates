from docx import Document
import comtypes.client

def replace_placeholders(caminho_arquivo, nome, cpf, caminho_saida):
    doc = Document(caminho_arquivo)

    def replace_in_runs(runs, placeholder, replacement):
        for run in runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)

    for par in doc.paragraphs:
        if "{NOME}" in par.text or "{CPF}" in par.text:
            replace_in_runs(par.runs, "{NOME}", nome)
            replace_in_runs(par.runs, "{CPF}", cpf)

    for table in doc.tables:
        for linha in table.rows:
            for celula in linha.cells:
                if "{NOME}" in celula.text or "{CPF}" in celula.text:
                    for par in celula.paragraphs:
                        replace_in_runs(par.runs, "{NOME}", nome)
                        replace_in_runs(par.runs, "{CPF}", cpf)

    doc.save(caminho_saida)
    
def save_as_pdf(caminho_docx, caminho_pdf):
    word = comtypes.client.CreateObject("Word.Application")
    word.Visible = False

    doc = word.Documents.Open(caminho_docx)

    doc.SaveAs(caminho_pdf, FileFormat=17)
    doc.Close()
    word.Quit()